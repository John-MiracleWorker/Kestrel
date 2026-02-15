"""
Document Intelligence Skill
Parse PDFs, Word documents (.docx), and Excel spreadsheets (.xlsx).
Extracts text content, metadata, and basic stats.

Dependencies: pypdf, python-docx, openpyxl (auto-installed on first use).
"""

import json
import logging
import os
import subprocess
import sys

logger = logging.getLogger("libre_bird.skills.documents")


def _ensure_dep(package: str, import_name: str = None):
    """Auto-install a dependency if missing."""
    import_name = import_name or package
    try:
        __import__(import_name)
    except ImportError:
        logger.info(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", package])


def _expand_path(path: str) -> str:
    """Expand ~ and resolve relative paths."""
    return os.path.expanduser(os.path.expandvars(path))


def tool_read_pdf(args: dict) -> dict:
    """Extract text from a PDF file."""
    path = _expand_path(args.get("path", ""))
    if not path:
        return {"error": "path is required"}
    if not os.path.exists(path):
        return {"error": f"File not found: {path}"}

    max_pages = int(args.get("max_pages", 20))

    _ensure_dep("pypdf")
    from pypdf import PdfReader

    try:
        reader = PdfReader(path)
        total_pages = len(reader.pages)
        pages_to_read = min(total_pages, max_pages)

        text_parts = []
        for i in range(pages_to_read):
            page_text = reader.pages[i].extract_text() or ""
            text_parts.append(f"--- Page {i+1} ---\n{page_text}")

        full_text = "\n\n".join(text_parts)

        # Truncate if too long
        if len(full_text) > 15000:
            full_text = full_text[:15000] + "\n\n... [truncated — showing first 15,000 chars]"

        metadata = {}
        if reader.metadata:
            for key in ["title", "author", "subject", "creator"]:
                val = getattr(reader.metadata, key, None)
                if val:
                    metadata[key] = str(val)

        return {
            "filename": os.path.basename(path),
            "total_pages": total_pages,
            "pages_read": pages_to_read,
            "metadata": metadata,
            "text": full_text,
        }
    except Exception as e:
        return {"error": str(e)}


def tool_read_docx(args: dict) -> dict:
    """Extract text from a Word document (.docx)."""
    path = _expand_path(args.get("path", ""))
    if not path:
        return {"error": "path is required"}
    if not os.path.exists(path):
        return {"error": f"File not found: {path}"}

    _ensure_dep("python-docx", "docx")

    try:
        from docx import Document
        doc = Document(path)

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Mark headings
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    text = f"{'#' * int(level) if level.isdigit() else '#'} {text}"
                paragraphs.append(text)

        # Also extract tables
        tables_text = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                tables_text.append(f"[Table {i+1}]\n" + "\n".join(rows))

        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n--- Tables ---\n" + "\n\n".join(tables_text)

        if len(full_text) > 15000:
            full_text = full_text[:15000] + "\n\n... [truncated]"

        # Core properties
        metadata = {}
        if doc.core_properties:
            for attr in ["author", "title", "subject", "created", "modified"]:
                val = getattr(doc.core_properties, attr, None)
                if val:
                    metadata[attr] = str(val)

        return {
            "filename": os.path.basename(path),
            "paragraphs": len(paragraphs),
            "tables": len(doc.tables),
            "metadata": metadata,
            "text": full_text,
        }
    except Exception as e:
        return {"error": str(e)}


def tool_read_xlsx(args: dict) -> dict:
    """Extract data from an Excel spreadsheet (.xlsx)."""
    path = _expand_path(args.get("path", ""))
    if not path:
        return {"error": "path is required"}
    if not os.path.exists(path):
        return {"error": f"File not found: {path}"}

    sheet_name = args.get("sheet", None)
    max_rows = int(args.get("max_rows", 100))

    _ensure_dep("openpyxl")

    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, read_only=True, data_only=True)

        result = {
            "filename": os.path.basename(path),
            "sheet_names": wb.sheetnames,
        }

        target_sheets = [sheet_name] if sheet_name else wb.sheetnames[:3]  # First 3 sheets
        sheets_data = {}

        for sname in target_sheets:
            if sname not in wb.sheetnames:
                continue
            ws = wb[sname]

            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= max_rows:
                    break
                row_data = [str(cell) if cell is not None else "" for cell in row]
                rows.append(row_data)

            sheets_data[sname] = {
                "rows_read": len(rows),
                "total_rows": ws.max_row or 0,
                "columns": ws.max_column or 0,
                "data": rows,
            }

        result["sheets"] = sheets_data
        wb.close()
        return result

    except Exception as e:
        return {"error": str(e)}


def tool_document_stats(args: dict) -> dict:
    """Get basic stats about a document without reading full content."""
    path = _expand_path(args.get("path", ""))
    if not path:
        return {"error": "path is required"}
    if not os.path.exists(path):
        return {"error": f"File not found: {path}"}

    ext = os.path.splitext(path)[1].lower()
    stats = {
        "filename": os.path.basename(path),
        "extension": ext,
        "size_bytes": os.path.getsize(path),
        "size_human": _human_size(os.path.getsize(path)),
    }

    try:
        if ext == ".pdf":
            _ensure_dep("pypdf")
            from pypdf import PdfReader
            reader = PdfReader(path)
            stats["pages"] = len(reader.pages)
            if reader.metadata and reader.metadata.title:
                stats["title"] = str(reader.metadata.title)
        elif ext == ".docx":
            _ensure_dep("python-docx", "docx")
            from docx import Document
            doc = Document(path)
            stats["paragraphs"] = len(doc.paragraphs)
            stats["tables"] = len(doc.tables)
        elif ext == ".xlsx":
            _ensure_dep("openpyxl")
            from openpyxl import load_workbook
            wb = load_workbook(path, read_only=True)
            stats["sheets"] = wb.sheetnames
            wb.close()
        else:
            stats["note"] = f"Unsupported format for detailed stats: {ext}"
    except Exception as e:
        stats["error"] = str(e)

    return stats


def _human_size(size: int) -> str:
    for unit in ["B", "KB", "MB", "GB"]:
        if size < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} TB"


# ---------------------------------------------------------------------------
# Exports
# ---------------------------------------------------------------------------

TOOL_DEFINITIONS = [
    {
        "type": "function",
        "function": {
            "name": "read_pdf",
            "description": "Extract text content, metadata, and page count from a PDF file.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the PDF file"},
                    "max_pages": {"type": "integer", "description": "Maximum pages to read (default 20)"},
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": "Extract text, headings, tables, and metadata from a Word document (.docx).",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the .docx file"},
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_xlsx",
            "description": "Extract data from an Excel spreadsheet (.xlsx). Returns rows as arrays.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the .xlsx file"},
                    "sheet": {"type": "string", "description": "Specific sheet name (optional, defaults to first 3)"},
                    "max_rows": {"type": "integer", "description": "Max rows to read per sheet (default 100)"},
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "document_stats",
            "description": "Get quick stats about a document (size, pages, sheets) without reading full content.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the document file"},
                },
                "required": ["path"],
            },
        },
    },
]

TOOL_HANDLERS = {
    "read_pdf": tool_read_pdf,
    "read_docx": tool_read_docx,
    "read_xlsx": tool_read_xlsx,
    "document_stats": tool_document_stats,
}
